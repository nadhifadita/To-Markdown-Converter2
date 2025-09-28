import fitz
from docx import Document
from collections import Counter
import argparse
import os
import sys

def extract_from_pdf(file_path):
    document_data = []
    try:
        document = fitz.open(file_path)
    except Exception as e:
        print(f"Error membuka file PDF: {e}")
        return []

    for page_num in range(len(document)):
        page = document.load_page(page_num)
        page_blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_SPANS)["blocks"]

        for block in page_blocks:
            if block["type"] == 0:
                for line in block["lines"]:
                    for span in line["spans"]:
                        text = span["text"].strip()
                        if text:
                            document_data.append({
                                "page_number": page_num + 1,
                                "text": text,
                                "size": round(span["size"], 2),
                                "font": span["font"],
                                "flags": span["flags"], 
                                "type": "text"
                            })
    document.close()
    return document_data

def extract_from_docx(file_path):
    document_data = []
    try:
        document = Document(file_path)
    except Exception as e:
        print(f"Error membuka file DOCX: {e}")
        return []

    for element in document.element.body:
        if element.tag.endswith('p'):
            para = document.paragraphs[element.getparent().index(element)]
            text_content = ""
            is_bold = False
            is_italic = False
            
            for run in para.runs:
                text_content += run.text
                if run.bold: is_bold = True
                if run.italic: is_italic = True

            document_data.append({
                "text": text_content.strip(),
                "style": para.style.name,
                "is_bold": is_bold,
                "is_italic": is_italic,
                "type": "paragraph"
            })
        
        elif element.tag.endswith('tbl'):
            table_data = []
            try:
                table = document.tables[element.getparent().index(element)]
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                document_data.append({
                    "data": table_data,
                    "type": "table"
                })
            except IndexError:
                continue

    return [item for item in document_data if item.get("text", "").strip() or item.get("data")]

def _group_pdf_elements(raw_data):
    if not raw_data: return []
    
    font_sizes = [d['size'] for d in raw_data if d['text']]
    normal_size = Counter(font_sizes).most_common(1)[0][0] if font_sizes else 10.0

    grouped_elements = []
    current_group = None

    for item in raw_data:
        text = item['text']
        if not text: continue
            
        is_same_style = (
            current_group and 
            current_group['size'] == item['size'] and 
            current_group['font'] == item['font'] and 
            (current_group['flags'] & 16) == (item['flags'] & 16)
        )
        
        is_heading_break = item['size'] > normal_size * 1.5 or "BAB" in text.upper()

        if current_group and is_same_style and not is_heading_break:
            current_group['text'] += " " + text
        else:
            if current_group:
                grouped_elements.append(current_group)
            
            current_group = item.copy()
            
    if current_group:
        grouped_elements.append(current_group)
        
    return grouped_elements

def parse_pdf_structure(raw_data):
    grouped_data = _group_pdf_elements(raw_data)
    final_structure = []
    
    if not grouped_data: return []
    
    font_sizes = [d['size'] for d in grouped_data if d['text']]
    unique_sorted_sizes = sorted(list(set(font_sizes)), reverse=True)
    size_map = {size: f"H{i+1}" for i, size in enumerate(unique_sorted_sizes[:3])}
    
    for item in grouped_data:
        text = item['text'].strip()
        if not text: continue
        
        level = "Paragraph"
        
        if item['size'] in size_map:
            level = size_map[item['size']]
            
        if text.upper().startswith(('BAB 1', 'BAB 2', 'BAB 3', 'BAB 4', 'BAB 5', 'BAB 6', 'BAB 7', 'DAFTAR')):
            level = "H1"
        elif text.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            if len(text.split('.', 2)) == 2:
                level = "H2"
            elif len(text.split('.', 3)) >= 3:
                level = "H3"
        elif text.strip().startswith(('-', '*')) or text.strip().endswith(('1', '2', '3')):
             level = "List_Item"

        final_structure.append({
            "text": text,
            "level": level,
            "type": "text_block",
            "is_bold": (item['flags'] & 16) > 0,
            "is_italic": (item['flags'] & 2) > 0 
        })

    return final_structure

def parse_docx_structure(raw_data):
    parsed_elements = []

    for item in raw_data:
        
        if item['type'] == 'table':
            parsed_elements.append(item)
            continue
            
        text = item['text'].strip()
        if not text: continue
        
        style = item.get('style', '').lower()
        level = "Paragraph"

        if style.startswith('heading 1'):
            level = "H1"
        elif style.startswith('heading 2'):
            level = "H2"
        elif style.startswith('heading 3'):
            level = "H3"
        elif style.startswith(('list paragraph', 'listbullet', 'listnumber')):
            level = "List_Item"
        elif text.startswith(('-', '*')) or text.startswith(('1.', '2.', '3.')):
             level = "List_Item"

        parsed_elements.append({
            "text": text,
            "level": level,
            "is_bold": item.get('is_bold', False),
            "is_italic": item.get('is_italic', False),
            "type": item['type']
        })

    return parsed_elements

def _apply_inline_formatting(text, is_bold, is_italic):
    if is_bold:
        text = f"**{text}**"
    if is_italic:
        if not is_bold:
             text = f"*{text}*"
    return text

def convert_to_markdown(parsed_structure):
    markdown_output = []

    for element in parsed_structure:
        text = element.get('text', '').strip()
        level = element.get('level')
        elem_type = element.get('type')
        
        if not text and elem_type != 'table':
            continue

        if elem_type in ['text_block', 'paragraph']:
            if level in ['Paragraph', 'List_Item']:
                 text = _apply_inline_formatting(
                    text, 
                    element.get('is_bold', False), 
                    element.get('is_italic', False)
                )

            if level == 'H1':
                markdown_output.append(f"# {text}")
            elif level == 'H2':
                markdown_output.append(f"## {text}")
            elif level == 'H3':
                markdown_output.append(f"### {text}")
            elif level == 'List_Item':
                if not text.strip().startswith(('-', '*', '1.', '2.', '3.')):
                    text = f"- {text}"
                markdown_output.append(text) 
            else: 
                markdown_output.append(text)

        elif elem_type == 'table':
            table_data = element.get('data', [])
            if table_data:
                header = table_data[0]
                num_cols = len(header)
                
                markdown_output.append("| " + " | ".join(header) + " |")
                separator = ["---"] * num_cols
                markdown_output.append("| " + " | ".join(separator) + " |")
                
                for row in table_data[1:]:
                    safe_row = (row + [''] * num_cols)[:num_cols]
                    markdown_output.append("| " + " | ".join(safe_row) + " |")

        markdown_output.append("")

    return "\n".join(markdown_output)

def get_file_extension(file_path):
    return os.path.splitext(file_path)[1].lower()

def run_conversion(input_path, output_path):
    if not os.path.exists(input_path):
        print(f"❌ Error: File input tidak ditemukan di jalur: {input_path}")
        return

    ext = get_file_extension(input_path)
    raw_data = None
    parser_func = None

    print(f"--- Memulai Konversi: {input_path} ---")
    print(f"1. Memulai Ekstraksi Teks (Tipe: {ext})...")

    if ext == '.pdf':
        raw_data = extract_from_pdf(input_path)
        parser_func = parse_pdf_structure
    elif ext == '.docx':
        raw_data = extract_from_docx(input_path)
        parser_func = parse_docx_structure
    else:
        print(f"❌ Error: Format file {ext} tidak didukung (Hanya PDF dan DOCX).")
        return

    if not raw_data:
        print("⚠️ Peringatan: Ekstraksi teks menghasilkan data kosong. Menghentikan proses.")
        return
        
    print(f"   -> Ekstraksi selesai. Ditemukan {len(raw_data)} elemen mentah.")
    
    print("2. Menganalisis dan Mendeteksi Struktur Dokumen...")
    parsed_structure = parser_func(raw_data)
    
    if not parsed_structure:
        print("⚠️ Peringatan: Parsing struktur menghasilkan data kosong. Menghentikan proses.")
        return
        
    print(f"   -> Deteksi selesai. Ditemukan {len(parsed_structure)} elemen terstruktur.")

    print("3. Mengkonversi Struktur ke Format Markdown...")
    markdown_content = convert_to_markdown(parsed_structure)
    print("   -> Konversi Markdown selesai.")

    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        print("\n✅ Proses Konversi Selesai!")
        print(f"   Output tersimpan di: {output_path}")
    except Exception as e:
        print(f"❌ Error saat menyimpan file output: {e}")

def main():
    parser = argparse.ArgumentParser(
        description="Alat Konversi Dokumen PDF/DOCX ke Markdown Terstruktur.",
        epilog="Contoh penggunaan: python converter_integrated.py input.pdf -o output.md"
    )
    
    parser.add_argument("input_file", help="Jalur ke file PDF atau DOCX.")
    parser.add_argument("-o", "--output", default=None, help="Nama file output.")
    
    args = parser.parse_args()
    
    if args.output is None:
        base_name = os.path.splitext(os.path.basename(args.input_file))[0]
        args.output = f"{base_name}_output.md"
    
    run_conversion(args.input_file, args.output)

if __name__ == '__main__':
    if len(sys.argv) == 1:
        input_file_path = 'contoh file skripsi.pdf' 
        output_file_path = 'skripsi_final_output.md'
        
        if os.path.exists(input_file_path):
             run_conversion(input_file_path, output_file_path)
        else:
             print("\n[PERHATIAN]: File 'contoh file skripsi.pdf' tidak ditemukan.")
             print("Silakan jalankan script ini dengan argumen CLI (e.g., python converter_integrated.py nama_file.pdf) atau letakkan file skripsi di direktori yang sama.")
    else:
        main()